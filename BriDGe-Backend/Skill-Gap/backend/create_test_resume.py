from docx import Document

doc = Document()
doc.add_heading('John Doe - Software Engineer', 0)

doc.add_heading('Skills', level=1)
doc.add_paragraph('Python, Java, machine learning, SQL, Docker')

doc.add_heading('Experience', level=1)
doc.add_paragraph('Worked as a Backend Developer building APIs in Python and writing SQL queries. Deployed with Docker.')

doc.save('test_resume.docx')
print("Created test_resume.docx")
